"""File Upload and Text Extraction Utility"""
import os
from pathlib import Path
from PIL import Image
import pytesseract
from docx import Document
import PyPDF2

# Configure Tesseract path for Windows
if os.name == 'nt':  # Windows
    tesseract_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Users\baves\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
    ]
    for tesseract_path in tesseract_paths:
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            break

# Try to import pdf2image for OCR on scanned PDFs
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
    
    # Configure Poppler path for Windows
    POPPLER_PATH = None
    if os.name == 'nt':
        poppler_paths = [
            r'C:\Users\baves\poppler\poppler-24.08.0\Library\bin',
            r'C:\Program Files\poppler\Library\bin',
            r'C:\Program Files\poppler\bin',
            os.path.expanduser(r'~\poppler\poppler-24.08.0\Library\bin'),
        ]
        for poppler_path in poppler_paths:
            if os.path.exists(poppler_path):
                POPPLER_PATH = poppler_path
                break
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    POPPLER_PATH = None

def extract_text_from_image(file_path):
    """Extract text from image using OCR"""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip() if text.strip() else "No text found in image (OCR returned empty result)"
    except Exception as e:
        return f"Error extracting text from image: {str(e)}"

def ocr_pdf_pages(file_path, max_pages=10):
    """Extract text from scanned PDF using OCR"""
    if not PDF2IMAGE_AVAILABLE:
        return None
    
    try:
        # Convert PDF pages to images (limit pages for performance)
        if POPPLER_PATH:
            images = convert_from_path(file_path, first_page=1, last_page=max_pages, dpi=150, poppler_path=POPPLER_PATH)
        else:
            images = convert_from_path(file_path, first_page=1, last_page=max_pages, dpi=150)
        
        text_parts = []
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image)
            if page_text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{page_text.strip()}")
        
        if text_parts:
            return "\n\n".join(text_parts)
        return None
    except Exception as e:
        print(f"OCR PDF error: {e}")
        return None

def extract_text_from_pdf(file_path):
    """Extract text from PDF file (with OCR fallback for scanned PDFs)"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            try:
                pdf_reader = PyPDF2.PdfReader(file)
            except Exception as e:
                error_str = str(e).lower()
                if 'password' in error_str or 'encrypt' in error_str:
                    return "Error: PDF is password-protected or encrypted. Please use an unprotected PDF."
                elif 'eof' in error_str or 'stream' in error_str:
                    return "Error: PDF file appears to be corrupted or incomplete."
                else:
                    return f"Error reading PDF: {str(e)}"
            
            # Limit to first 50 pages to avoid timeout
            max_pages = min(len(pdf_reader.pages), 50)
            for i in range(max_pages):
                try:
                    page_text = pdf_reader.pages[i].extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_error:
                    # Skip problematic pages
                    continue
            
            if max_pages < len(pdf_reader.pages):
                text += f"\n[Note: Only first {max_pages} of {len(pdf_reader.pages)} pages extracted]"
        
        # If no text found, try OCR
        if not text.strip():
            ocr_text = ocr_pdf_pages(file_path)
            if ocr_text:
                return ocr_text + "\n\n[Extracted via OCR from scanned PDF]"
            else:
                if PDF2IMAGE_AVAILABLE:
                    return "No text found in PDF. OCR also failed - the PDF may be empty or unreadable."
                else:
                    return "No text found in PDF (scanned/image-based). Install 'pdf2image' and Poppler for OCR support."
        
        return text.strip()
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        return text.strip() if text.strip() else "No text found in document"
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        # Try multiple encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                    text = file.read().strip()
                    if text:
                        return text
            except:
                continue
        return "No text found in file"
    except Exception as e:
        return f"Error reading text file: {str(e)}"

def extract_text_from_eml(file_path):
    """Extract text from email (.eml) file"""
    try:
        import email
        from email import policy
        from email.parser import BytesParser
        
        with open(file_path, 'rb') as file:
            msg = BytesParser(policy=policy.default).parse(file)
        
        text_parts = []
        
        # Get subject
        if msg['subject']:
            text_parts.append(f"Subject: {msg['subject']}")
        
        # Get from
        if msg['from']:
            text_parts.append(f"From: {msg['from']}")
        
        # Get body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    body = part.get_content()
                    if body:
                        text_parts.append(body)
        else:
            body = msg.get_content()
            if body:
                text_parts.append(body)
        
        text = "\n".join(text_parts)
        return text.strip() if text.strip() else "No text found in email file"
    except Exception as e:
        # Fallback to reading as text
        return extract_text_from_txt(file_path)

def extract_text_from_file(file_path):
    """
    Extract text from various file types
    Supports: Images (JPG, PNG), PDF, DOCX, TXT, EML
    """
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
        return extract_text_from_image(file_path)
    elif file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        return extract_text_from_docx(file_path)
    elif file_ext == '.doc':
        # .doc files are not supported by python-docx, try reading as binary text
        return "Error: Old .doc format not supported. Please convert to .docx or .txt"
    elif file_ext == '.txt':
        return extract_text_from_txt(file_path)
    elif file_ext == '.eml':
        return extract_text_from_eml(file_path)
    else:
        return f"Unsupported file type: {file_ext}"
